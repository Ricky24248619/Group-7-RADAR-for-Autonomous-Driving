"""Check the Markdown list cases used by the generated handbooks."""

import pathlib
import sys
import unittest

try:
    from docx import Document
except ImportError:
    Document = None

if Document is not None:
    sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1] / "scripts"))
    from build_handbook import convert


@unittest.skipIf(Document is None, "python-docx is an optional document-building dependency")
class HandbookTests(unittest.TestCase):
    def test_wrapped_items_and_separate_numbered_lists(self):
        doc = Document()
        convert("1. First **item**\n   continued here\n2. Second\n\n"
                "A paragraph.\n\n1. Restart\n\n- Bullet\n  continuation\n\n"
                "| Key | Value |\n|---|---|\n| A | B |", doc)
        paras = [p for p in doc.paragraphs if p.text]
        self.assertEqual([p.text for p in paras], [
            "First item continued here", "Second", "A paragraph.",
            "Restart", "Bullet continuation",
        ])
        self.assertEqual(paras[-1].style.name, "List Bullet")
        starts = []
        for para in (paras[0], paras[1], paras[3]):
            num_id = para._p.pPr.numPr.numId.val
            num = doc.part.numbering_part.element.num_having_numId(num_id)
            starts.append(int(num.xpath("w:lvlOverride/w:startOverride/@w:val")[0]))
        self.assertEqual(starts, [1, 2, 1])
        self.assertEqual(doc.tables[0].cell(1, 1).text, "B")


if __name__ == "__main__":
    unittest.main()
