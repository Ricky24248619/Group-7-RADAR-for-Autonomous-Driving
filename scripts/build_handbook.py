#!/usr/bin/env python3
"""Generate PROJECT-HANDBOOK.docx from PROJECT-HANDBOOK.md.

The Markdown is the source of truth; the Word file exists so the handbook can be shared
with people who do not read the repository — the client, the auditor, next year's team.
Regenerating rather than maintaining two copies is what stops them drifting apart.

    pip install python-docx
    python scripts/build_handbook.py

Handles the subset of Markdown the handbook actually uses: headings, paragraphs, pipe
tables, fenced code, bullet and numbered lists, blockquotes, and inline bold/italic/code.
"""

import pathlib
import re
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

REPO = pathlib.Path(__file__).resolve().parent.parent
SOURCE = REPO / "docs" / "PROJECT-HANDBOOK.md"
OUTPUT = REPO / "docs" / "PROJECT-HANDBOOK.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(paragraph, text):
    """Render inline **bold**, *italic* and `code` into a paragraph."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            # Markdown links render as "text (url)"; nothing here needs live hyperlinks.
            paragraph.add_run(re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", part))


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], [r for r in cells[2:]]   # row 1 is the |---| separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        add_runs(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells_out = table.add_row().cells
        for i, text in enumerate(row[:len(header)]):
            cells_out[i].text = ""
            add_runs(cells_out[i].paragraphs[0], text)
    doc.add_paragraph()


def add_code(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert(md, doc):
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or set(stripped) == {"-"} and len(stripped) >= 3:
            i += 1
            continue

        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = doc.add_heading(level=min(level, 4))
            add_runs(heading, stripped[level:].strip())
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(para, " ".join(b for b in block if b))
            for run in para.runs:
                run.italic = True
            continue

        if re.match(r"^[-*] ", stripped):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            add_runs(doc.add_paragraph(style="List Number"),
                     re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        # Otherwise a paragraph: join until a blank line or a new block starts.
        block = []
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^\s*([#>|`]|[-*] |\d+\. )", lines[i]):
            block.append(lines[i].strip())
            i += 1
        if block:
            add_runs(doc.add_paragraph(), " ".join(block))
        else:
            i += 1


def main():
    if not SOURCE.is_file():
        sys.exit(f"Source not found: {SOURCE}")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    convert(SOURCE.read_text(encoding="utf-8"), doc)
    doc.save(OUTPUT)

    words = len(SOURCE.read_text(encoding="utf-8").split())
    print(f"Wrote {OUTPUT.relative_to(REPO)}")
    print(f"  from {SOURCE.relative_to(REPO)} ({words:,} words)")
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")


if __name__ == "__main__":
    main()
